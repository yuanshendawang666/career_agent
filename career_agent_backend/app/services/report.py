"""
报告生成服务模块

提供职业发展报告的数据准备、文档生成、文本润色、学习资源推荐等功能。
"""
import json
import os
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.llm_client import call_qwen
from app.database.mysql import SessionLocal
from app.models.job import JobTitleProfile
from app.models.region_stats import JobRegionStats
from app.models.student import Student
from app.services.graph import get_job_paths
from app.services.matching import compute_match

plt.rcParams["font.sans-serif"] = ["SimHei"]
plt.rcParams["axes.unicode_minus"] = False


# ========== Word 样式辅助函数 ==========
def set_paragraph_background(paragraph, color_hex):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def set_paragraph_left_border(paragraph, color_hex, size=12):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def set_paragraph_top_border(paragraph, color_hex, size=6):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size))
    top.set(qn('w:color'), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)


def set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_card_paragraph(doc, text, bg_color_hex, left_border_color_hex=None, top_border_color_hex=None):
    p = doc.add_paragraph()
    set_paragraph_background(p, bg_color_hex)
    if left_border_color_hex:
        set_paragraph_left_border(p, left_border_color_hex)
    if top_border_color_hex:
        set_paragraph_top_border(p, top_border_color_hex)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.size = Pt(12)
    return p


def add_section_heading(doc, text, level=1):
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(8)
    return heading


def remove_table_borders(table):
    """去除表格所有边框"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)


# ========== 核心业务函数 ==========
def get_report_data(db: Session, student_id: int, job_title: str) -> Dict[str, Any]:
    student = db.query(Student).filter(Student.id == student_id).first()
    if not student:
        raise ValueError("学生不存在")
    student_profile = student.profile_json

    job_profile = db.query(JobTitleProfile).filter(JobTitleProfile.job_title == job_title).first()
    if not job_profile:
        raise ValueError(f"岗位 {job_title} 的典型画像不存在")

    match_details = compute_match(student_profile, job_profile)
    paths = get_job_paths(job_title)

    region_stats = db.query(JobRegionStats).filter(JobRegionStats.job_title == job_title).all()
    region_stats_data = []
    for stat in region_stats:
        top_cities = []
        if stat.top_cities:
            try:
                top_cities = json.loads(stat.top_cities)
            except json.JSONDecodeError:
                top_cities = []
        region_stats_data.append({
            "region": stat.region,
            "demand_count": stat.demand_count,
            "salary_min_avg": float(stat.salary_min_avg) if stat.salary_min_avg else 0,
            "salary_max_avg": float(stat.salary_max_avg) if stat.salary_max_avg else 0,
            "top_cities": top_cities,
        })

    action_plan = generate_action_plan(student_profile, job_profile, match_details)
    transition_advice = generate_transition_advice(student_profile, job_profile, paths)
    learning_resources = generate_learning_resources(student_profile, job_profile, match_details)
    gap_text = generate_gap_analysis(student_profile, job_profile, match_details)

    return {
        "student": {
            'id': student.id,
            "name": student.name,
            "education": student_profile.get("education", ""),
            "major": student_profile.get("major", ""),
            "skills": student_profile.get("skills", []),
            "certificates": student_profile.get("certificates", []),
            "overall_score": student_profile.get("overall_score", 0),
            "overall_reason": student_profile.get("overall_reason", ""),
        },
        "job_title": job_title,
        "match_details": match_details,
        "paths": paths,
        "region_stats": region_stats_data,
        "gap_analysis": gap_text,
        "action_plan": action_plan,
        "evaluation_cycle": "建议每3-6个月进行一次自我评估，根据学习进度和行业变化调整计划。",
        "transition_advice": transition_advice,
        "learning_resources": learning_resources
    }


def generate_report(db: Session, student_id: int, job_title: str) -> str:
    data = get_report_data(db, student_id, job_title)
    return generate_report_from_data(data)


def generate_report_from_data(data: Dict[str, Any]) -> str:
    """完整版导出：卡片样式 + 雷达图 + 页眉页脚（稳定版）"""
    try:
        student = data["student"]
        job_title = data["job_title"]
        match_details = data["match_details"]
        paths = data["paths"]
        region_stats = data["region_stats"]
        gap_analysis = data["gap_analysis"]
        action_plan = data["action_plan"]
        evaluation_cycle = data.get("evaluation_cycle", "建议每3-6个月进行一次自我评估，根据学习进度和行业变化调整计划。")
        transition_advice = data.get("transition_advice", "")
        learning_resources = data.get("learning_resources", "")

        now = datetime.now()
        # 精确到秒的日期时间
        date_str = now.strftime("%Y%m%d_%H%M%S")
        # 过滤非法字符
        safe_name = re.sub(r'[\\/*?:"<>|]', '', student['name'])
        safe_job = re.sub(r'[\\/*?:"<>|]', '', job_title)
        # 最终文件名：用户名_岗位名称_职业发展报告_日期时间.docx
        filename = f"{safe_name}_{safe_job}_职业发展报告_{date_str}.docx"

        doc = Document()
        # 窄页边距
        for section in doc.sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

        # ===================== 页面边距 + 页眉页脚间距（缩小版）=====================
        section = doc.sections[0]
        # 缩小页眉距离顶部
        section.header_distance = Cm(0.4)
        # 缩小页脚距离底部
        section.footer_distance = Cm(0.3)

        # ===================== 页眉（修复版）=====================
        header = section.header
        try:
            header_para = header.paragraphs[0]
        except:
            header_para = header.add_paragraph()
        header_para.clear()

        # 左右布局
        header_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_ALIGN_PARAGRAPH.RIGHT)

        run_left = header_para.add_run("「智途」— 你的私人职业规划助手")
        run_left.font.name = 'Times New Roman'
        run_left._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_left.font.size = Pt(10)

        header_para.add_run("\t")
        run_right = header_para.add_run(now.strftime("%Y-%m-%d %H:%M:%S"))
        run_right.font.name = 'Times New Roman'
        run_right._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_right.font.size = Pt(10)

        # 默认正文字体
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        style.font.size = Pt(12)

        # 主标题
        title = doc.add_heading(f"职业生涯发展报告 - {student['name']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.size = Pt(24)
            run.font.bold = True

        # 一、学生基本信息
        add_section_heading(doc, "一、学生基本信息")
        info_text = (
            f"姓名：{student['name']}\n"
            f"学历：{student['education']}\n"
            f"专业：{student['major']}\n"
            f"技能：{', '.join(student['skills'])}\n"
            f"证书：{', '.join(student['certificates'])}\n"
            f"综合竞争力评分：{student['overall_score']}分\n"
            f"评分理由：{student['overall_reason']}"
        )
        add_card_paragraph(doc, info_text, 'E0F2FE')

        # 二、人岗匹配分析
        add_section_heading(doc, "二、人岗匹配分析")
        summary_text = f"目标岗位：{job_title}\n综合匹配度：{match_details['total_score']:.1%}"
        add_card_paragraph(doc, summary_text, 'FFEDD5')

        # 维度表格
        dimensions = [
            ("基础要求", "base_match"), ("职业技能", "professional_match"),
            ("职业素养", "quality_match"), ("发展潜力", "potential_match"),
            ("技能匹配", "skill_match"), ("证书匹配", "cert_match"),
            ("创新匹配", "innovation_match"), ("学习匹配", "learning_match"),
            ("抗压匹配", "stress_match"), ("沟通匹配", "communication_match"),
            ("学历匹配", "education_match"), ("专业匹配", "major_match"),
            ("经验匹配", "experience_match"), ("语言匹配", "language_match"),
        ]
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "维度"
        hdr_cells[1].text = "匹配度"
        for label, key in dimensions:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = f"{match_details.get(key, 0):.1%}"

        # 雷达图
        try:
            img_buffer = generate_radar_chart(match_details)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_buffer, width=Inches(5))
        except Exception as e:
            print(f"雷达图生成失败: {e}")

        # 三、职业目标与路径规划
        add_section_heading(doc, "三、职业目标与路径规划")
        add_card_paragraph(doc, f"推荐职业目标：{job_title}", 'DBEAFE', left_border_color_hex='3B82F6')
        if paths.get("promotions"):
            add_card_paragraph(doc, f"晋升路径：{' → '.join(paths['promotions'])}", 'DCFCE7', left_border_color_hex='10B981')
        if paths.get("transfers"):
            add_card_paragraph(doc, f"横向换岗路径：{'、'.join(paths['transfers'])}", 'FEF3C7', left_border_color_hex='F59E0B')

        # 四、区域机会分析
        if region_stats:
            add_section_heading(doc, "四、区域机会分析")
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "区域"
            hdr[1].text = "招聘数量"
            hdr[2].text = "平均薪资(K)"
            hdr[3].text = "主要城市"
            for idx, stat in enumerate(region_stats):
                row = table.add_row().cells
                row[0].text = stat["region"]
                row[1].text = str(stat["demand_count"])
                row[2].text = f"{stat['salary_min_avg']:.1f} - {stat['salary_max_avg']:.1f}K"
                row[3].text = ", ".join(stat["top_cities"])
                bg_color = ['FFF7ED', 'F0FDFA', 'FAF5FF', 'FDF2F8'][idx % 4]
                for cell in row:
                    set_cell_background(cell, bg_color)

        # 五、差距分析
        add_section_heading(doc, "五、差距分析")
        add_card_paragraph(doc, gap_analysis, 'F1F5F9')

        # 六、PDCA发展计划
        add_section_heading(doc, "六、PDCA发展计划")
        sections = {}
        current = None
        for line in action_plan.split('\n'):
            line = line.strip()
            if re.match(r'^(Plan|Do|Check|Act)[：:]', line, re.I):
                current = line.split('：')[0] if '：' in line else line.split(':')[0]
                sections[current] = []
            elif current and line:
                sections[current].append(line)
        if not sections:
            sections = {'Plan': [action_plan]}
        pdca_colors = {
            'Plan': ('F0F9FF', '3B82F6'),
            'Do': ('F0FDF4', '10B981'),
            'Check': ('FEFCE8', 'F59E0B'),
            'Act': ('F5F3FF', '8B5CF6'),
        }
        for key in ['Plan', 'Do', 'Check', 'Act']:
            if key in sections:
                content = '\n'.join(sections[key])
                bg_color, border_color = pdca_colors[key]
                add_card_paragraph(doc, f"{key}：\n{content}", bg_color, left_border_color_hex=border_color)

        # 七、岗位调动建议
        if transition_advice:
            add_section_heading(doc, "七、岗位调动建议")
            items = re.split(r'(?=\d+\.\s+)', transition_advice)
            colors = [('E0F2FE', '38BDF8'), ('DCFCE7', '4ADE80'), ('FEF3C7', 'FBBF24')]
            for idx, item in enumerate(items):
                if not item.strip():
                    continue
                match = re.match(r'(\d+\.)\s*(.*)', item.strip(), re.DOTALL)
                if match:
                    number, content = match.groups()
                    bg_color, border_color = colors[idx % 3]
                    add_card_paragraph(doc, f"{number} {content}", bg_color, left_border_color_hex=border_color)
                else:
                    add_card_paragraph(doc, item.strip(), 'F8FAFC')

        # 八、学习资源推荐
        if learning_resources:
            add_section_heading(doc, "八、学习资源推荐")
            stages = re.split(r'(?=【)', learning_resources)
            resource_colors = [
                ('FFE4E6', 'F43F5E'), ('CFFAFE', '06B6D4'), ('ECFCCB', '84CC16'), ('FAE8FF', 'D946EF')
            ]
            for idx, stage in enumerate(stages):
                if not stage.strip():
                    continue
                match = re.match(r'【([^】]+)】([\s\S]*)', stage.strip())
                if match:
                    title, content = match.groups()
                    duration = ''
                    dur_match = re.search(r'(建议用时|用时)\s*([0-9~\-]+)\s*个月', title)
                    if dur_match:
                        duration = f"{dur_match.group(1)} {dur_match.group(2)} 个月"
                        title = title.replace(dur_match.group(0), '').strip()
                    bg_color, border_color = resource_colors[idx % 4]
                    p = doc.add_paragraph()
                    set_paragraph_background(p, bg_color)
                    set_paragraph_top_border(p, border_color, size=6)
                    run = p.add_run(f"{title}")
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                    run.font.bold = True
                    run.font.size = Pt(14)
                    if duration:
                        run2 = p.add_run(f"  {duration}")
                        run2.font.name = 'Times New Roman'
                        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                        run2.italic = True
                    p.add_run(f"\n{content}")
                    p.paragraph_format.space_after = Pt(12)
                else:
                    add_card_paragraph(doc, stage.strip(), 'F8FAFC')

        # 九、评估周期
        add_section_heading(doc, "九、评估周期与动态调整")
        add_card_paragraph(doc, evaluation_cycle, 'F1F5F9')

        # 结尾提示
        doc.add_paragraph()
        tip_para = doc.add_paragraph()
        tip_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tip_run = tip_para.add_run("内容由AI生成，请仔细甄别")
        tip_run.font.name = 'Times New Roman'
        tip_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        tip_run.font.size = Pt(10)
        tip_run.font.italic = True
        tip_run.font.color.rgb = RGBColor(128, 128, 128)

        os.makedirs(settings.report_dir, exist_ok=True)
        file_path = os.path.join(settings.report_dir, filename)
        doc.save(file_path)
        return file_path
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise


# ---------- AI 生成函数 ----------
def generate_action_plan(student_profile, job_profile, match_details) -> str:
    prompt = f"""
你是一位职业规划导师。请为学生制定一份详细的职业发展行动计划，严格按照 **PDCA（计划-执行-检查-调整）** 框架组织内容。输出要求：
- 不使用任何 Markdown 格式符号（如 **、*、-、--- 等）。
- 用普通文字标题标识每个部分，例如："Plan："、"Do："、"Check："、"Act："。
- 列表项使用数字序号（1. 2. 3.）或圆点（•），但不要使用短横线（-）或星号（*）。
- 每部分内容清晰分段，条理明确。

示例格式（纯文本）：
Plan：
1. 职业目标：获得XX公司offer
2. 学术/技能目标：掌握XX技术
3. 平台/资源：参与XX社区

Do：
1. 具体行动1：学习XX课程
2. 具体行动2：投递XX公司
3. 具体行动3：每周XX小时实践

Check：
1. 评估周期：每月复盘
2. 评估指标：技能测试成绩、项目进度
3. 偏差识别：分析滞后原因

Act：
1. 备选方案：若A计划受阻，转向B计划
2. 资源补充：寻求导师帮助

现在，请根据以下学生能力和目标岗位，制定符合该格式的PDCA计划：

学生当前能力：
- 技能：{', '.join(student_profile.get('skills', []))}
- 证书：{', '.join(student_profile.get('certificates', []))}
- 创新能力评分：{student_profile.get('innovation_score', 0)}/5
- 学习能力评分：{student_profile.get('learning_score', 0)}/5
- 抗压能力评分：{student_profile.get('stress_score', 0)}/5
- 沟通能力评分：{student_profile.get('communication_score', 0)}/5
- 学历：{student_profile.get('education', '未提供')}
- 专业：{student_profile.get('major', '未提供')}
- 实习经历：{', '.join(student_profile.get('internships', []))}

目标岗位：{job_profile.job_title}
岗位要求：
- 技能：{', '.join(json.loads(job_profile.skills) if isinstance(job_profile.skills, str) else job_profile.skills)}
- 证书：{', '.join(json.loads(job_profile.certificates) if isinstance(job_profile.certificates, str) else job_profile.certificates)}
- 学历要求：{job_profile.education_required or '无明确要求'}
- 工作经验要求：{job_profile.experience_required or '无明确要求'}
- 其他要求：{job_profile.other_requirements or '无'}

当前匹配度：{match_details['total_score']:.1%}

请严格按照上述纯文本格式输出，不要包含任何 Markdown 符号。
"""
    result = call_qwen(
        user_prompt=prompt,
        system_prompt="你是一名职业规划导师，擅长为学生制定个性化的成长计划。请输出纯文本，不要使用 Markdown 格式。",
        max_tokens=1200,
        temperature=0.5
    )
    result = re.sub(r'\*\*|__', '', result)
    return result


def generate_gap_analysis(
    student_profile: Dict[str, Any],
    job_profile: JobTitleProfile,
    match_details: Dict[str, float],
) -> str:
    gaps = []
    if match_details.get("skill_match", 0) < 0.5:
        gaps.append("技能匹配度较低，需补充目标岗位的核心技能。")
    if match_details.get("cert_match", 0) < 0.5:
        gaps.append("证书方面存在差距，建议考取相关认证。")
    if match_details.get("education_match", 0) < 0.5:
        gaps.append("学历要求不匹配，可考虑提升学历或通过项目经验弥补。")
    if match_details.get("internship_match", 0) < 0.5:
        gaps.append("缺乏相关实习经验，建议寻找对口实习机会。")
    if not gaps:
        gaps.append("学生能力与岗位要求基本匹配，但仍需在细节上进一步优化。")
    return " ".join(gaps)


def generate_radar_chart(match_details: Dict[str, float]) -> BytesIO:
    dimension_map = {
        "base_match": "基础要求",
        "professional_match": "职业技能",
        "quality_match": "职业素养",
        "potential_match": "发展潜力",
        "skill_match": "技能匹配",
        "cert_match": "证书匹配",
        "innovation_match": "创新匹配",
        "learning_match": "学习匹配",
        "stress_match": "抗压匹配",
        "communication_match": "沟通匹配",
        "education_match": "学历匹配",
        "major_match": "专业匹配",
        "experience_match": "经验匹配",
        "language_match": "语言匹配",
    }

    labels = []
    values = []
    for key, label in dimension_map.items():
        val = match_details.get(key, 0.0)
        if val > 0:
            labels.append(label)
            values.append(val)

    if not labels:
        fig, ax = plt.subplots(figsize=(8, 8))
        ax.text(0.5, 0.5, '暂无匹配数据', ha='center', va='center', fontsize=20)
        ax.axis('off')
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=120, bbox_inches='tight')
        plt.close(fig)
        img_buffer.seek(0)
        return img_buffer

    values += values[:1]
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(9, 9), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color="blue", alpha=0.25)
    ax.plot(angles, values, color="blue", linewidth=2)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=20)
    for tick in ax.get_xticklabels():
        tick.set_rotation(30)
        tick.set_horizontalalignment('right')

    ax.set_ylim(0, 1)
    ax.set_yticks([0.2, 0.4, 0.6, 0.8, 1.0])
    ax.set_yticklabels(["20%", "40%", "60%", "80%", "100%"], fontsize=16)
    ax.grid(True)

    img_buffer = BytesIO()
    plt.savefig(img_buffer, format="png", dpi=120, bbox_inches="tight")
    plt.close(fig)
    img_buffer.seek(0)
    return img_buffer


def generate_learning_resources(
    student_profile: Dict[str, Any],
    job_profile: JobTitleProfile,
    match_details: Dict[str, float],
) -> str:
    gaps = []
    if match_details.get("skill_match", 0) < 0.5:
        gaps.append("技能")
    if match_details.get("cert_match", 0) < 0.5:
        gaps.append("证书")
    if match_details.get("education_match", 0) < 0.5:
        gaps.append("学历")
    if match_details.get("internship_match", 0) < 0.5:
        gaps.append("实习经历")

    if not gaps:
        return "您的能力与岗位要求基本匹配，建议通过项目实践进一步提升综合能力。"

    skills_str = ", ".join(student_profile.get("skills", []))
    certs_str = ", ".join(student_profile.get("certificates", []))
    internships_str = ", ".join(student_profile.get("internships", []))

    job_skills = job_profile.skills
    if isinstance(job_skills, str):
        try:
            job_skills = json.loads(job_skills)
        except json.JSONDecodeError:
            job_skills = []
    job_skills_str = ", ".join(job_skills)

    job_certs = job_profile.certificates
    if isinstance(job_certs, str):
        try:
            job_certs = json.loads(job_certs)
        except json.JSONDecodeError:
            job_certs = []
    job_certs_str = ", ".join(job_certs)

    system_prompt = "你是一位职业规划导师，擅长为学生制定学习计划。请输出纯文本，不要使用 Markdown 格式（如 **、[]() 等）。"
    user_prompt = f"""
学生当前能力：
- 技能：{skills_str}
- 证书：{certs_str}
- 实习经历：{internships_str}

目标岗位：{job_profile.job_title}
岗位要求：
- 技能：{job_skills_str}
- 证书：{job_certs_str}
- 学历：{job_profile.education_required}
- 工作经验：{job_profile.experience_required}

主要差距：{', '.join(gaps)}

请为学生制定一份详细的学习计划，按阶段（如第一阶段、第二阶段、第三阶段）列出每个阶段应该学习的具体内容（细化到知识点或技能），并给出合理的时间安排（例如每周投入多少小时，预计完成周期）。输出格式要求：
- 使用纯文本，不要包含任何 Markdown 符号（如 **、*、-、---、[]() 链接等）。
- 每个阶段标题用【第一阶段：阶段名称】表示。
- 阶段内容用 • 或数字序号列出。
- 不要输出任何网址或超链接。

直接返回纯文本计划，使用以下格式：
【第一阶段：阶段名称】建议用时 X-X 个月
• 具体内容1
• 具体内容2
...
"""
    try:
        result = call_qwen(
            user_prompt=user_prompt,
            system_prompt=system_prompt,
            max_tokens=1200,
            temperature=0.5,
        )
        result = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', result)
        result = re.sub(r'\*\*', '', result)
        return result.strip()
    except Exception as e:
        print(f"生成学习计划失败: {e}")
        return ""


def generate_transition_advice(
    student_profile: Dict[str, Any],
    job_profile: JobTitleProfile,
    paths: Dict[str, List[str]],
) -> str:
    promotions = paths.get("promotions", [])
    transfers = paths.get("transfers", [])
    if not promotions and not transfers:
        return "暂无明确的职业发展路径数据。"

    skills_str = ", ".join(student_profile.get("skills", []))
    certs_str = ", ".join(student_profile.get("certificates", []))
    internships_str = ", ".join(student_profile.get("internships", []))
    education = student_profile.get("education", "")
    major = student_profile.get("major", "")
    work_exp = student_profile.get("work_experience", "")
    language = student_profile.get("language", "")

    job_skills = job_profile.skills
    if isinstance(job_skills, str):
        try:
            job_skills = json.loads(job_skills)
        except json.JSONDecodeError:
            job_skills = []
    job_skills_str = ", ".join(job_skills)

    job_certs = job_profile.certificates
    if isinstance(job_certs, str):
        try:
            job_certs = json.loads(job_certs)
        except json.JSONDecodeError:
            job_certs = []
    job_certs_str = ", ".join(job_certs)

    prompt = f"""
学生当前能力：
- 技能：{skills_str}
- 证书：{certs_str}
- 实习经历：{internships_str}
- 学历：{education}
- 专业：{major}
- 工作经验：{work_exp}
- 语言能力：{language}

目标岗位：{job_profile.job_title}
岗位要求：
- 技能：{job_skills_str}
- 证书：{job_certs_str}
- 学历：{job_profile.education_required}
- 工作经验：{job_profile.experience_required}

可能的职业发展路径：
- 晋升路径：{' → '.join(promotions) if promotions else '暂无'}
- 换岗路径：{'、'.join(transfers) if transfers else '暂无'}

请为学生提供如何实现这些职业发展的具体建议，包括：
1. 如何晋升到更高岗位（需要学习哪些技能、考取哪些证书、积累哪些经验）。
2. 如何转换到相关岗位（需要补充哪些技能、参与哪些项目、拓展哪些人脉）。

输出要求：
- 纯文本格式，不要使用任何 Markdown 符号（如 **、*、-、--- 等）。
- 用数字序号（1. 2. 3.）或圆点（•）列出要点。
- 每部分建议清晰分段。

请用流畅的段落文字描述，分点说明。
"""
    try:
        result = call_qwen(
            user_prompt=prompt,
            system_prompt="你是一位职业发展导师，擅长为学生提供职业转型和晋升的具体建议。请输出纯文本，不要使用 Markdown 格式。",
            max_tokens=800,
            temperature=0.5,
        )
        result = re.sub(r'\*\*', '', result)
        return result
    except Exception as e:
        print(f"生成岗位调动建议失败: {e}")
        return "暂无法生成具体的岗位调动建议，请参考职业路径自行规划。"


def polish_text(text: str, instruction: Optional[str] = None) -> str:
    if not text or len(text.strip()) < 10:
        return text
    if instruction:
        user_prompt = f"请根据以下要求润色文本：{instruction}\n\n原文：{text}\n\n请只输出润色后的文本，不要输出任何额外解释、不要重复要求。"
    else:
        user_prompt = f"请润色以下文本，使其更通顺、专业，不要改变原意。只输出润色后的文本，不要输出其他内容。\n\n{text}"
    system_prompt = "你是一位专业的文字编辑，擅长优化语言表达。"
    result = call_qwen(user_prompt, system_prompt, max_tokens=500, temperature=0.5)
    return result.strip()


def add_hyperlinks_to_paragraph(paragraph, text: str) -> None:
    pattern = r"\[([^\]]+)\]\(([^)]+)\)"
    last_index = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_index:
            paragraph.add_run(text[last_index: match.start()])
        link_text = match.group(1)
        url = match.group(2)
        run = paragraph.add_run(link_text)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True
        last_index = match.end()
    if last_index < len(text):
        paragraph.add_run(text[last_index:])


def generate_path_report(data: dict) -> str:
    student = data['student']
    path_name = data['path_name']
    match_score = data.get('match_score', 0)
    advice = data.get('advice', '')
    alternative_path = data.get('alternative_path', '')
    development_plan = data.get('development_plan', '')

    doc = Document()
    doc.add_heading(f'职业规划报告 - {student["name"]}', 0)

    doc.add_heading('一、学生基本信息', level=1)
    doc.add_paragraph(f'姓名：{student["name"]}')
    doc.add_paragraph(f'专业：{student.get("major", "未提供")}')
    doc.add_paragraph(f'年级：{student.get("grade", "未提供")}')
    doc.add_paragraph(f'意向城市：{student.get("intended_city", "未提供")}')

    doc.add_heading('二、当前能力分析', level=1)
    doc.add_paragraph(f'技能：{", ".join(student.get("skills", []))}')
    doc.add_paragraph(f'证书：{", ".join(student.get("certificates", []))}')
    doc.add_paragraph(f'综合竞争力评分：{student.get("overall_score", 0)}分')
    doc.add_paragraph(f'评分理由：{student.get("overall_reason", "")}')

    doc.add_heading('三、目标路径与匹配度', level=1)
    doc.add_paragraph(f'推荐职业路径：{path_name}')
    doc.add_paragraph(f'当前匹配度：{match_score}%')
    if advice:
        doc.add_paragraph(f'改进建议：{advice}')
    if alternative_path:
        doc.add_paragraph(f'备选路径：{alternative_path}')

    doc.add_heading('四、PDCA 发展计划', level=1)
    doc.add_paragraph(development_plan)

    doc.add_heading('五、评估与调整', level=1)
    doc.add_paragraph('建议每3个月更新一次个人信息，系统将自动重新评估匹配度并调整计划。')

    os.makedirs(settings.report_dir, exist_ok=True)
    safe_path_name = "".join(c for c in path_name if c.isalnum() or c in (' ', '-', '_')).replace(' ', '_')
    filename = f"path_report_{student['id']}_{safe_path_name}.docx"
    file_path = os.path.join(settings.report_dir, filename)
    doc.save(file_path)
    return file_path

